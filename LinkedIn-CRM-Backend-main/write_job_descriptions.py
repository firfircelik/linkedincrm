import openai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

import openai
import time
from openai import OpenAI
# Set up your OpenAI API credentials
client = OpenAI(
    # defaults to os.environ.get("OPENAI_API_KEY")
    api_key="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",
)


def format_job_description(file_path, jobtitle):
    doc = Document(file_path)
    # Set a uniform font and size
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Add header and footer
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Job Description"
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Confidential - For internal use only"
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Adjust headings and add bullet points
    for paragraph in doc.paragraphs:
        if 'Job Title:' in paragraph.text or 'Responsibilities:' in paragraph.text or \
           'Experience/ Skills required:' in paragraph.text or 'Qualifications:' in paragraph.text:
            paragraph.style = doc.styles['Heading 1']
        
        if paragraph.text.startswith('- '):
            paragraph.style = doc.styles['List Bullet']
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Save the formatted document
    if jobtitle:
        formatted_file_path = jobtitle + ".docx"
    else:
        formatted_file_path = "jd.docx"
    doc.save(formatted_file_path)
    return formatted_file_path

def generate_detailed_job_description(profile_name, experience, location):
    # Assuming 'client' is previously defined as an OpenAI client


    prompt = f"""Create a detailed and structured job description for a lead who has an experience in {experience} and lead location is {location}. The job description should follow the continuity of the original offer but in more detail and structured in a professional way that you would expect to see in an official job description. The candidate should see their career reflected in this job offer so that they’re more enticed to engage due to them feeling like they meet the requirements.
                The sections within the job description should Include:
                Job Title: 


                What is the job title for this position? Consider the job title that targeted candidates will search for.
                You may wish to consider the level of person you are recruiting and the impact you want the job title to have. This alone can have a significant impact on engagement and applications - so make it enticing to the candidate in question.
                This job should be perceived as a step up from their current role.
                If they are a founder or have lots of experience, then it might be worth opting for a consultancy role in your offer as that would seem more attractive to someone at that stage in their career. 


                Location:


                When choosing the location you’ll need to analyse the candidates profile and determine what this person's ideal location would be. If they seem firmly rooted and settled in their current  location then choose this location that they’re already in or a neighbouring city so that they would perceive this as being a convenient transition.  
                If their profile suggests that they have an extensive expatriate career, then it would also be worth opting for another country entirely. This analysis is for your discernment from analysing their profile.
                The location should also state some flexibility. When generating this consider:
                Will travel be required?
                Is the position home based?
                Will relocation be an option? 


                Employment type:


                This should be based on the analysis taken to generate the ideal job for this candidate. Consider questions like:
                What is the employment type? Permanent, Contract, Full Time, Part time? 
                And decide what this candidate would most likely engage with at this stage in their career. 


                Overview of company:


                This is your chance to present your company and create a great first impression. 
                You need a cohesive and positive message about the company and the client you are representing and put personality into it.
                Consider what is special about this company; provide detail about the history, share the vision and what makes them attractive to potential employees. 
                                        
                            
                Purpose of position 
                                    
                Use this as an opportunity to pitch the position. By making this section compelling you will increase candidate engagement, receive more applicants and recruit better staff. 	


                Industry:
                                        
                Providing industry information will make your job description more appealing. 					
                Specialisation:
                Pick skill sets and any other attributes that you have harvested from the candidate's profile and try to include that here so that the candidate will feel that the job is meant for them due to them having specific skills required for this role.
                Responsibilities:
                What are the main responsibilities of the position? This is a great chance to sell the position and provide clear details on what is required. 
                Experience/ skills required:
                What experience is required for this position? Be clear on essential and desirable skills.
                When generating the following subjects, do not include anything that is outside of the candidates current level of expertise. 				
                Some important areas for consideration are:
                                            
                Years of experience - match the candidates tenure of experience
                                            
                Type of experience (Technical / Sales / Marketing / Accounting
                experience) - Match the candidates current experience
                                            
                Function / sector experience 
                                            
                Skill specific to position
                                            
                Management experience (state this as optional)
                                            
                What personable skills are required (well organised, self-
                                                
                motivated, articulate, team player, leader, dedication, outward
                                                
                going, multi taking skills etc.)
                                            
                Computer Skills
                                            
                Languages (stick to only English as a requirement)
                Salary:
                By offering information, if competitive, you will increase the quality and quantity of applicants. 
                When choosing the salary, make it very competitive to the highest industry standard that you would likely see in the role on offer. Must include exact salary(do not include range) in local currency or dollars.
                Salary cont. (OTE + benefits):
                Choose typical corporate benefits as well as high industry standard to endorse the roles benefits more. 
                The salary should be in dollars or the local currency that the job offer is in.
                The salary range should be very attractive for this person. 
                Qualifications:
                Make this match the candidates current qualifications and if they don’t have any stated on their profile mention that the qualifications aren’t mandatory and experience is rather required
                Eligibility to work:
                Mention visa sponsorship is available in the country that is stated
                Mandatory Values:
                In section headings only Include section name ends with :"	
                Do not include any “how to apply” section 	
                Do not include company name section and do not include any company name in whole content"""
    



    #prompt = f"Create a detailed and structured job description for a person with {experience} years of experience and candidate location is {location}. The job poition should be upgraded from current position. Include only sections for Job Title, Job Type, Roles, Responsibilities, Skills, Location and Salary Range. The salary range should be very attractive for this person and in euros. Do not include How to Apply section. In section headings only Include section name ends with :"

    chat_completion = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="gpt-4",
    )
    print(chat_completion.choices[0].message.content)
    time.sleep(10)  # Pause to simulate processing time

    job_description = chat_completion.choices[0].message.content

    # Create a Word document
    doc = Document()
    
    #job_description_heading = doc.add_heading('Job Description', 0)

    #for run in job_description_heading.runs:
    #    run.font.size = Pt(24)  # Increase the font size as needed

    # Process the structured output
    sections = job_description.split('\n\n')
    job_title = ""
    for section in sections:
        if 'Job Title:' in section:
            # Extract job title from the section
            job_title = section.split(':', 1)[1].strip()
        if ':' in section:
            header, content = section.split(':', 1)
            doc.add_heading(header.strip() + ":", level=2)
            # Create a new paragraph for content, indented under the header
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Pt(90)  # Adjust the indentation as necessary
            paragraph.add_run(content.strip())
        else:
            # Handle sections without a clear header
            # Create a new paragraph for content, indented under the header
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Pt(90)  # Adjust the indentation as necessary
            paragraph.add_run(section.strip())

        # Set font size and paragraph formatting for all content
        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                run.font.size = Pt(12)
     # Save the document
    filename = f"{profile_name.replace(' ', '_')}_Job_Description.docx"
    doc.save(filename)
    filename = format_job_description(filename, job_title)

    return filename


# Example usage
generate_detailed_job_description('John', 'Marketing Specialist with a background in digital marketing', 'Melbourne, Australia')
